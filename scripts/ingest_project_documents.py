#!/usr/bin/env python3
"""Build Mongo-ready JSONL files for Astor Butler project documents."""

from __future__ import annotations

import argparse
import hashlib
import json
import mimetypes
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from pypdf import PdfReader


@dataclass(frozen=True)
class SourceDocument:
    path: Path
    category: str
    title: str
    tags: tuple[str, ...]


SOURCES = [
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_Butler_Final_With_Header.docx"), "diploma", "Astor Butler final diploma with HSE title", ("diploma", "hse", "main")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_Butler_Triada_69_with_sources.docx"), "diploma", "Astor Butler triad with sources", ("diploma", "triad", "sources")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_Butler_Academic_Format.docx"), "diploma", "Astor Butler academic format", ("diploma", "academic")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Titulnik_Astor_Butler_2025.docx"), "diploma", "Astor Butler title page 2025", ("diploma", "title")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_Butler_Triada_69_with_TOC.docx"), "diploma", "Astor Butler triad with TOC", ("diploma", "triad", "toc")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_Butler_Triada_69_ready_for_TOC.docx"), "diploma", "Astor Butler triad ready for TOC", ("diploma", "triad", "toc")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/astor_butler_pitch_deck_may_2025_final.pdf"), "presentation", "Astor Butler pitch deck May 2025 final", ("pitch", "presentation")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/astor_butler_pitch_deck_may_2025.pptx.pdf"), "presentation", "Astor Butler pitch deck May 2025 PPTX PDF", ("pitch", "presentation")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_butler_one_pager(ru).docx"), "one-pager", "Astor Butler one pager RU DOCX", ("one-pager", "ru")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_butler_one_pager(ru).pdf"), "one-pager", "Astor Butler one pager RU PDF", ("one-pager", "ru")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/astor_butler_speech.txt"), "speech", "Astor Butler speech", ("speech",)),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor_Butler_Conference_Paper.docx"), "paper", "Astor Butler conference paper", ("conference", "paper")),
    SourceDocument(Path("/Users/michaelwelly/Downloads/Astor Butler Conf.docx"), "paper", "Astor Butler conference draft", ("conference", "paper")),
    SourceDocument(Path("/Users/michaelwelly/IdeaProjects/Astor_Butler_MVP/docs/ARCHITECTURE.md"), "architecture", "Astor Butler architecture", ("architecture", "repo")),
    SourceDocument(Path("/Users/michaelwelly/IdeaProjects/Astor_Butler_MVP/README.md"), "architecture", "Astor Butler README", ("readme", "repo")),
    SourceDocument(Path("/Users/michaelwelly/Obsidian/Astor_Butler_Knowledge/01_Project/Project_Context.md"), "memory", "Project context", ("obsidian", "context")),
    SourceDocument(Path("/Users/michaelwelly/Obsidian/Astor_Butler_Knowledge/01_Project/Work_Plan.md"), "memory", "Work plan", ("obsidian", "plan")),
    SourceDocument(Path("/Users/michaelwelly/Obsidian/Astor_Butler_Knowledge/04_Tech/Tech_Decisions.md"), "memory", "Tech decisions", ("obsidian", "tech")),
    SourceDocument(Path("/Users/michaelwelly/Obsidian/Astor_Butler_Knowledge/03_FSM/FSM_Index.md"), "memory", "FSM index", ("obsidian", "fsm")),
]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out-dir", default="/private/tmp/astor_document_ingest")
    parser.add_argument("--chunk-size", type=int, default=2400)
    args = parser.parse_args()

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    docs_path = out_dir / "project_documents.jsonl"
    chunks_path = out_dir / "document_chunks.jsonl"
    run_path = out_dir / "document_ingest_runs.jsonl"
    summary_path = out_dir / "summary.md"

    now = datetime.now(timezone.utc).isoformat()
    run_id = hashlib.sha256(now.encode("utf-8")).hexdigest()[:16]
    documents = []
    chunks = []
    missing = []

    for source in SOURCES:
        if not source.path.exists():
            missing.append(str(source.path))
            continue
        text = extract_text(source.path)
        doc_id = stable_id(source.path)
        stat = source.path.stat()
        document = {
            "_id": doc_id,
            "runId": run_id,
            "title": source.title,
            "category": source.category,
            "tags": list(source.tags),
            "sourcePath": str(source.path),
            "extension": source.path.suffix.lower(),
            "mimeType": mimetypes.guess_type(source.path.name)[0] or "application/octet-stream",
            "sha256": sha256(source.path),
            "sizeBytes": stat.st_size,
            "modifiedAt": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
            "ingestedAt": now,
            "textLength": len(text),
            "wordCount": len(text.split()),
            "preview": normalize_space(text)[:800],
            "text": text,
        }
        documents.append(document)
        for index, chunk_text in enumerate(chunk_texts(text, args.chunk_size)):
            chunks.append({
                "_id": f"{doc_id}:{index:04d}",
                "documentId": doc_id,
                "runId": run_id,
                "chunkIndex": index,
                "sourcePath": str(source.path),
                "category": source.category,
                "tags": list(source.tags),
                "text": chunk_text,
                "textLength": len(chunk_text),
            })

    write_jsonl(docs_path, documents)
    write_jsonl(chunks_path, chunks)
    write_jsonl(run_path, [{
        "_id": run_id,
        "startedAt": now,
        "documentCount": len(documents),
        "chunkCount": len(chunks),
        "missing": missing,
    }])
    summary_path.write_text(summary(documents, chunks, missing, run_id), encoding="utf-8")
    print(summary(documents, chunks, missing, run_id))
    print(f"JSONL_DIR={out_dir}")


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    return ""


def extract_docx(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def chunk_texts(text: str, chunk_size: int) -> Iterable[str]:
    normalized = normalize_space(text)
    if not normalized:
        return
    start = 0
    while start < len(normalized):
        yield normalized[start:start + chunk_size]
        start += chunk_size


def normalize_space(text: str) -> str:
    return " ".join(text.split())


def stable_id(path: Path) -> str:
    return hashlib.sha256(str(path).encode("utf-8")).hexdigest()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_jsonl(path: Path, rows: list[dict]) -> None:
    with path.open("w", encoding="utf-8") as file:
        for row in rows:
            file.write(json.dumps(row, ensure_ascii=False) + "\n")


def summary(documents: list[dict], chunks: list[dict], missing: list[str], run_id: str) -> str:
    lines = [
        f"# Astor Butler Document Ingest {run_id}",
        "",
        f"- documents: {len(documents)}",
        f"- chunks: {len(chunks)}",
        f"- missing: {len(missing)}",
        "",
        "## Documents",
    ]
    for document in documents:
        lines.append(f"- {document['title']} ({document['category']}): {document['wordCount']} words, {document['sourcePath']}")
    if missing:
        lines.extend(["", "## Missing"])
        lines.extend(f"- {item}" for item in missing)
    return "\n".join(lines)


if __name__ == "__main__":
    main()
